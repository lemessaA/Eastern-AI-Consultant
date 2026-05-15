"""Document parsing utilities (PDF, DOCX, TXT, CSV)."""

from __future__ import annotations

import csv
import io
from pathlib import Path

import pypdf
from docx import Document as DocxDocument

from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_text(file_path: str | Path, mime_type: str | None = None) -> str:
    """Best-effort plain-text extraction from a local file."""
    path = Path(file_path)
    suffix = path.suffix.lower().lstrip(".")
    mime = (mime_type or "").lower()

    try:
        if suffix == "pdf" or "pdf" in mime:
            return _extract_pdf(path)
        if suffix in {"docx", "doc"} or "word" in mime:
            return _extract_docx(path)
        if suffix == "csv" or "csv" in mime:
            return _extract_csv(path)
        # Fallback: read as utf-8 text.
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        logger.warning("document_extract_failed", path=str(path), error=str(exc))
        return ""


def _extract_pdf(path: Path) -> str:
    out: list[str] = []
    with path.open("rb") as fh:
        reader = pypdf.PdfReader(fh)
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                out.append(text)
    return "\n\n".join(out)


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_csv(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as fh:
        reader = csv.reader(fh)
        for row in reader:
            rows.append(", ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def extract_text_from_bytes(data: bytes, mime_type: str, filename: str = "upload") -> str:
    """Extract text without writing to disk first."""
    mime = mime_type.lower()
    try:
        if "pdf" in mime or filename.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(io.BytesIO(data))
            return "\n\n".join((p.extract_text() or "") for p in reader.pages)
        if "word" in mime or filename.lower().endswith((".doc", ".docx")):
            doc = DocxDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return data.decode("utf-8", errors="ignore")
    except Exception as exc:
        logger.warning("document_bytes_extract_failed", error=str(exc))
        return ""
